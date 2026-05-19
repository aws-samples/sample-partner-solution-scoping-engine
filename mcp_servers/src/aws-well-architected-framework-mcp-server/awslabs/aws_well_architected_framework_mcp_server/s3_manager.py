"""S3 management module for WAFR report storage."""

import json
import logging
import os
import sys
import tempfile
from datetime import datetime
from typing import Any, Dict, Optional
import re
from io import BytesIO

import boto3
from botocore.exceptions import ClientError

# NEW: Using unified DOCX generator (SOW approach)
from .docx_generator import generate_wafr_docx

logger = logging.getLogger(__name__)

def clean_text_symbols(text: str) -> str:
    """Remove unknown symbols and emojis from text."""
    # Remove emoji and special symbols
    text = re.sub(r'[^\x00-\x7F]+', '', text)  # Remove non-ASCII characters
    
    # Replace common problematic symbols with clean alternatives
    replacements = {
        'Tools': 'Tools',
        'Security': 'Security',
        'Shield': 'Shield', 
        'Performance': 'Performance',
        'Cost': 'Cost',
        'Sustainability': 'Sustainability',
        'Critical': 'Critical',
        'Warning': 'Warning',
        'Info': 'Info',
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Clean up any remaining problematic characters and keep basic punctuation
    text = re.sub(r'[^\w\s\-_.,!?()[\]{}:;"\'/\\@#$%^&*+=<>|~`•—–]', '', text)
    
    return text

def markdown_to_docx(markdown_content: str) -> bytes:
    """Convert markdown content to DOCX format (fallback method)."""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Clean the markdown content first
        clean_content = clean_text_symbols(markdown_content)
        
        doc = Document()
        
        # Split content into lines
        lines = clean_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Handle headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('---'):
                # Add page break for section separators
                doc.add_page_break()
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('• ') or line.startswith('- '):
                # Bullet points
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # Regular paragraph
                if line:
                    doc.add_paragraph(line)
        
        # Save to bytes
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer.getvalue()
        
    except ImportError:
        logger.error("python-docx not installed, falling back to cleaned markdown")
        return clean_text_symbols(markdown_content).encode('utf-8')
    except Exception as e:
        logger.error(f"Error converting to DOCX: {e}")
        return clean_text_symbols(markdown_content).encode('utf-8')

def get_s3_config() -> Dict[str, str]:
    """Get S3 configuration from SERA CustomerConfig or environment fallback."""
    from .consts import get_aws_region
    
    try:
        # Try to import SERA CustomerConfig first
        current_dir = os.path.dirname(os.path.abspath(__file__))
        backend_path = os.path.abspath(os.path.join(current_dir, '..', '..', '..', '..', '..', 'backend'))
        if backend_path not in sys.path:
            sys.path.insert(0, backend_path)
        
        from config.app_config import CustomerConfig
        
        # Load configuration if not already loaded
        if not CustomerConfig._config:
            CustomerConfig.load_config()
        
        bucket = CustomerConfig.get_sow_s3_bucket()
        region = CustomerConfig.get_aws_region()
        
        logger.info(f"S3 Config loaded: bucket={bucket}, region={region}")
        
        return {
            "bucket": bucket,
            "region": region,
        }
    except Exception as e:
        logger.error(f"Failed to load CustomerConfig: {e}")
        # Fallback to environment variables using centralized get_aws_region()
        bucket = os.getenv("S3_UPLOAD_BUCKET")
        region = get_aws_region()
        logger.info(f"Using fallback S3 config: bucket={bucket}, region={region}")
        return {
            "bucket": bucket,
            "region": region,
        }

def create_s3_client():
    """Create S3 client with proper configuration."""
    config = get_s3_config()
    return boto3.client('s3', region_name=config["region"])

async def upload_wafr_docx_to_s3(chat_id: str, file_path: str) -> Dict[str, Any]:
    """Upload WAFR DOCX document to S3 with versioning and metadata (like SOW server).
    
    Args:
        chat_id: Unique chat identifier
        file_path: Local path to the DOCX file
        
    Returns:
        Dictionary containing S3 upload results
    """
    try:
        config = get_s3_config()
        s3_client = create_s3_client()
        
        # Construct S3 key following SERA pattern: {chatId}/wafr/wafr_assessment.docx
        s3_key = f"{chat_id}/wafr/wafr_assessment.docx"
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Prepare metadata
        metadata = {
            "chat-id": chat_id,
            "document-type": "wafr-assessment",
            "generation-timestamp": datetime.now().isoformat(),
            "file-size": str(file_size),
            "generated-by": "aws-well-architected-framework-mcp-server"
        }
        
        logger.info(f"Uploading WAFR DOCX to S3: s3://{config['bucket']}/{s3_key}")
        
        # Upload file to S3
        with open(file_path, 'rb') as f:
            response = s3_client.put_object(
                Bucket=config["bucket"],
                Key=s3_key,
                Body=f,
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                Metadata=metadata,
                ServerSideEncryption="AES256"
            )
        
        # Generate presigned URL for download (valid for 1 hour)
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': config["bucket"], 'Key': s3_key},
            ExpiresIn=3600
        )
        
        # Save metadata to S3
        metadata_key = f"{chat_id}/wafr/wafr_assessment.metadata.json"
        metadata_content = {
            "chat_id": chat_id,
            "document_type": "wafr-assessment",
            "s3_key": s3_key,
            "file_size": file_size,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "generation_timestamp": datetime.now().isoformat(),
            "format": "docx",
            "professional_formatting": True
        }
        
        s3_client.put_object(
            Bucket=config["bucket"],
            Key=metadata_key,
            Body=json.dumps(metadata_content, indent=2),
            ContentType="application/json",
        )
        
        # Match cost-analysis return structure exactly
        s3_url = f"s3://{config['bucket']}/{s3_key}"
        
        return {
            "s3_url": s3_url,
            "s3_key": s3_key,
            "version_id": response.get("VersionId"),
            "file_size": file_size,
            "bucket": config["bucket"],
            "metadata_key": metadata_key,
            "filename": "wafr_assessment",
            "format": "docx"
        }
        
    except Exception as e:
        logger.error(f"Failed to upload WAFR DOCX to S3: {e}")
        raise Exception(f"Failed to upload WAFR DOCX to S3: {str(e)}")

async def upload_wafr_report_to_s3(chat_id: str, report_content: str, format: str = 'markdown') -> Dict[str, Any]:
    """Upload WAFR report to S3 with professional formatting.
    
    Args:
        chat_id: Unique chat identifier
        report_content: The report content as string
        format: Report format ('markdown' or 'docx')
        
    Returns:
        Dictionary containing S3 upload results
    """
    try:
        config = get_s3_config()
        s3_client = create_s3_client()
        
        # Construct S3 key following SERA pattern: {chatId}/wafr/wafr_assessment.{ext}
        file_extension = 'docx' if format == 'docx' else 'md'
        s3_key = f"{chat_id}/wafr/wafr_assessment.{file_extension}"
        
        # Get content size and prepare content
        if format == 'docx':
            # Use professional DOCX generation
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                docx_path = temp_docx.name
            
            try:
                await generate_wafr_docx(report_content, docx_path)
                
                # Read the generated DOCX file
                with open(docx_path, 'rb') as f:
                    content_bytes = f.read()
                
                # Clean up temp file
                os.unlink(docx_path)
                
            except Exception as e:
                logger.error(f"Professional DOCX generation failed, falling back to basic: {e}")
                # Fallback to basic DOCX generation
                content_bytes = markdown_to_docx(report_content)
            
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        elif format == 'markdown':
            content_bytes = clean_text_symbols(report_content).encode('utf-8')
            content_type = "text/markdown"
        else:
            # For other formats, report_content would be bytes
            content_bytes = report_content if isinstance(report_content, bytes) else report_content.encode('utf-8')
            content_type = "application/octet-stream"
        
        content_size = len(content_bytes)
        
        # Prepare metadata
        metadata = {
            "chat-id": chat_id,
            "document-type": "wafr-assessment",
            "generation-timestamp": datetime.now().isoformat(),
            "file-size": str(content_size),
            "generated-by": "aws-well-architected-framework-mcp-server",
            "format": format,
            "professional-formatting": "true" if format == 'docx' else "false"
        }
        
        logger.info(f"Uploading professional WAFR report to S3: s3://{config['bucket']}/{s3_key}")
        
        # Upload content to S3
        response = s3_client.put_object(
            Bucket=config["bucket"],
            Key=s3_key,
            Body=content_bytes,
            ContentType=content_type,
            Metadata=metadata,
            ServerSideEncryption="AES256"
        )
        
        # Generate presigned URL for download (valid for 1 hour)
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': config["bucket"], 'Key': s3_key},
            ExpiresIn=3600
        )
        
        # Save metadata to S3
        metadata_key = f"{chat_id}/wafr/wafr_assessment.metadata.json"
        metadata_content = {
            "chat_id": chat_id,
            "document_type": "wafr-assessment",
            "s3_key": s3_key,
            "file_size": content_size,
            "content_type": content_type,
            "generation_timestamp": datetime.now().isoformat(),
            "format": format,
            "professional_formatting": format == 'docx'
        }
        
        s3_client.put_object(
            Bucket=config["bucket"],
            Key=metadata_key,
            Body=json.dumps(metadata_content, indent=2),
            ContentType="application/json",
        )
        
        return {
            "success": True,
            "bucket": config["bucket"],
            "s3_key": s3_key,
            "s3_url": f"s3://{config['bucket']}/{s3_key}",
            "presigned_url": presigned_url,
            "file_size": content_size,
            "content_type": content_type,
            "etag": response.get("ETag", "").strip('"'),
            "version_id": response.get("VersionId"),
            "metadata_key": metadata_key,
            "professional_formatting": format == 'docx'
        }
        
    except Exception as e:
        logger.error(f"Failed to upload WAFR report to S3: {e}")
        return {"success": False, "error": str(e)}

async def retrieve_diagram_data(chat_id: str) -> Dict[str, Any]:
    """Enhanced diagram data retrieval with comprehensive service extraction."""
    try:
        config = get_s3_config()
        s3_client = create_s3_client()
        
        # Try to get diagram metadata
        metadata_key = f"{chat_id}/diagram/diagram.metadata.json"
        try:
            response = s3_client.get_object(Bucket=config["bucket"], Key=metadata_key)
            metadata = json.loads(response['Body'].read().decode('utf-8'))
            
            # Enhanced service extraction from diagram metadata
            enhanced_services = extract_services_from_diagram_metadata(metadata)
            
            return {
                "success": True,
                "metadata": metadata,
                "enhanced_services": enhanced_services,
                "diagram_url": f"s3://{config['bucket']}/{chat_id}/diagram/diagram.png"
            }
        except ClientError:
            return {"success": False, "error": "No diagram data found"}
            
    except Exception as e:
        logger.error(f"Error retrieving diagram data: {e}")
        return {"success": False, "error": str(e)}

def extract_services_from_diagram_metadata(metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Extract comprehensive AWS services and configurations from diagram metadata."""
    services = []
    configurations = {}
    
    # Extract from existing aws_services if available
    if metadata.get("aws_services"):
        services.extend(metadata["aws_services"])
    
    # Extract from diagram code if available
    diagram_code = metadata.get("diagram_code", "")
    if diagram_code:
        # Enhanced service extraction from diagram code
        code_services = extract_services_from_code(diagram_code)
        services.extend(code_services)
        
        # Extract configurations from code
        configurations.update(extract_configurations_from_code(diagram_code))
    
    # Extract from generation metadata
    if metadata.get("generation_metadata"):
        gen_meta = metadata["generation_metadata"]
        if gen_meta.get("services_used"):
            services.extend(gen_meta["services_used"])
    
    return {
        "aws_services": list(set(services)),
        "configurations": configurations,
        "metadata_source": "diagram"
    }

def extract_services_from_code(code: str) -> list:
    """Extract AWS services from diagram generation code."""
    services = []
    
    # Common AWS service patterns in diagram code
    service_patterns = {
        'EC2': ['EC2', 'Instance'],
        'S3': ['S3', 'Bucket'],
        'RDS': ['RDS', 'Database', 'MySQL', 'PostgreSQL'],
        'Lambda': ['Lambda', 'Function'],
        'API Gateway': ['APIGateway', 'RestApi'],
        'CloudFront': ['CloudFront', 'Distribution'],
        'Route 53': ['Route53', 'DNS'],
        'ELB': ['ELB', 'LoadBalancer'],
        'ALB': ['ALB', 'ApplicationLoadBalancer'],
        'VPC': ['VPC', 'Network'],
        'CloudWatch': ['CloudWatch', 'Monitoring'],
        'IAM': ['IAM', 'Role', 'Policy'],
        'ElastiCache': ['ElastiCache', 'Redis', 'Memcached'],
        'SNS': ['SNS', 'Topic'],
        'SQS': ['SQS', 'Queue'],
        'DynamoDB': ['DynamoDB', 'Table'],
        'KMS': ['KMS', 'Key'],
        'WAF': ['WAF', 'WebApplicationFirewall'],
        'Shield': ['Shield', 'DdosProtection'],
        'Cognito': ['Cognito', 'UserPool'],
        'Systems Manager': ['SystemsManager', 'ParameterStore'],
        'Auto Scaling': ['AutoScaling', 'ScalingGroup'],
        'Certificate Manager': ['CertificateManager', 'Certificate'],
        'AWS Backup': ['Backup', 'BackupVault']
    }
    
    code_upper = code.upper()
    for service_name, patterns in service_patterns.items():
        for pattern in patterns:
            if pattern.upper() in code_upper:
                services.append(service_name)
                break
    
    return services

def extract_configurations_from_code(code: str) -> Dict[str, Any]:
    """Extract configuration details from diagram code."""
    configurations = {}
    
    # Multi-AZ detection
    if any(pattern in code.lower() for pattern in ['multi-az', 'multiple availability', 'cross-az']):
        configurations["multi_az"] = True
    
    # Backup detection
    if any(pattern in code.lower() for pattern in ['backup', 'snapshot', 'recovery']):
        configurations["backup_configured"] = True
    
    # Caching detection
    if any(pattern in code.lower() for pattern in ['cache', 'redis', 'memcached', 'cloudfront']):
        configurations["caching_configured"] = True
    
    # Cross-region detection
    if any(pattern in code.lower() for pattern in ['cross-region', 'replica', 'disaster recovery']):
        configurations["cross_region_configured"] = True
    
    # Monitoring detection
    if any(pattern in code.lower() for pattern in ['cloudwatch', 'monitoring', 'metrics']):
        configurations["monitoring_configured"] = True
    
    return configurations

async def retrieve_cost_analysis_data(chat_id: str) -> Dict[str, Any]:
    """Enhanced cost analysis data retrieval."""
    try:
        config = get_s3_config()
        s3_client = create_s3_client()
        
        # Try to get cost analysis report
        cost_key = f"{chat_id}/cost-analysis/cost_analysis.md"
        try:
            response = s3_client.get_object(Bucket=config["bucket"], Key=cost_key)
            cost_content = response['Body'].read().decode('utf-8')
            
            # Try to get cost metadata
            metadata_key = f"{chat_id}/cost-analysis/cost_analysis.metadata.json"
            metadata = {}
            try:
                meta_response = s3_client.get_object(Bucket=config["bucket"], Key=metadata_key)
                metadata = json.loads(meta_response['Body'].read().decode('utf-8'))
            except ClientError:
                pass
            
            # Enhanced cost insights extraction
            enhanced_insights = extract_enhanced_cost_insights(cost_content, metadata)
                
            return {
                "success": True,
                "cost_report": cost_content,
                "metadata": metadata,
                "enhanced_insights": enhanced_insights,
                "cost_url": f"s3://{config['bucket']}/{cost_key}"
            }
        except ClientError:
            return {"success": False, "error": "No cost analysis data found"}
            
    except Exception as e:
        logger.error(f"Error retrieving cost analysis data: {e}")
        return {"success": False, "error": str(e)}

async def retrieve_cloudformation_template(chat_id: str) -> Dict[str, Any]:
    """Retrieve CloudFormation template from S3."""
    try:
        config = get_s3_config()
        s3_client = create_s3_client()
        
        # Try to get CloudFormation template
        cfn_key = f"{chat_id}/cloudformation/template.yaml"
        try:
            response = s3_client.get_object(Bucket=config["bucket"], Key=cfn_key)
            cfn_content = response['Body'].read().decode('utf-8')
            
            return {
                "success": True,
                "template": cfn_content,
                "template_url": f"s3://{config['bucket']}/{cfn_key}"
            }
        except ClientError:
            return {"success": False, "error": "No CloudFormation template found"}
            
    except Exception as e:
        logger.error(f"Error retrieving CloudFormation template: {e}")
        return {"success": False, "error": str(e)}

def extract_enhanced_cost_insights(cost_content: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Extract enhanced cost insights from cost analysis content and metadata."""
    insights = {}
    
    import re
    
    # Extract total monthly cost with better patterns
    cost_patterns = [
        r'total.*?cost.*?\$(\d+(?:,\d+)*(?:\.\d+)?)',
        r'monthly.*?cost.*?\$(\d+(?:,\d+)*(?:\.\d+)?)',
        r'\$(\d+(?:,\d+)*(?:\.\d+)?).*?per month',
        r'\$(\d+(?:,\d+)*(?:\.\d+)?).*?monthly'
    ]
    
    for pattern in cost_patterns:
        match = re.search(pattern, cost_content, re.IGNORECASE)
        if match:
            insights["monthly_cost"] = match.group(1)
            break
    
    # Enhanced optimization opportunities detection
    optimization_keywords = {
        'reserved_instances_opportunity': ['reserved instance', 'ri savings', 'on-demand premium'],
        'right_sizing_opportunity': ['right-siz', 'oversized', 'underutilized', 'cpu utilization'],
        'storage_optimization': ['storage class', 'infrequent access', 'glacier'],
        'data_transfer_optimization': ['data transfer', 'bandwidth', 'cloudfront'],
        'unused_resources': ['unused', 'idle', 'stopped instances']
    }
    
    for opportunity, keywords in optimization_keywords.items():
        if any(keyword in cost_content.lower() for keyword in keywords):
            insights[opportunity] = True
    
    # Extract service cost breakdown with enhanced patterns
    service_cost_patterns = [
        r'(\w+).*?\$(\d+(?:,\d+)*(?:\.\d+)?)',
        r'(\w+)\s+service.*?\$(\d+(?:,\d+)*(?:\.\d+)?)',
        r'(\w+)\s+costs.*?\$(\d+(?:,\d+)*(?:\.\d+)?)'
    ]
    
    service_costs = {}
    for pattern in service_cost_patterns:
        matches = re.findall(pattern, cost_content, re.IGNORECASE)
        for service, cost in matches:
            if service.upper() in ['EC2', 'S3', 'RDS', 'LAMBDA', 'CLOUDFRONT']:
                service_costs[service.upper()] = cost
    
    if service_costs:
        insights["service_breakdown"] = service_costs
    
    # Extract from metadata if available
    if metadata.get("cost_breakdown"):
        insights["detailed_breakdown"] = metadata["cost_breakdown"]
    
    return insights
