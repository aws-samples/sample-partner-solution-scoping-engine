"""
Professional Document Formatter for WAFR Reports.

This module provides professional DOCX document formatting with proper styling,
structure, and visual elements for enterprise-grade WAFR assessment reports.
"""

import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


logger = logging.getLogger(__name__)


class ProfessionalDocumentFormatter:
    """
    Creates professionally formatted DOCX documents for WAFR reports.
    
    Provides consistent styling, proper structure, and visual elements
    for enterprise-grade assessment reports.
    """
    
    # Color scheme (AWS/Professional)
    COLOR_GREEN = RGBColor(40, 167, 69)  # #28a745
    COLOR_YELLOW = RGBColor(255, 193, 7)  # #ffc107
    COLOR_RED = RGBColor(220, 53, 69)  # #dc3545
    COLOR_BLUE = RGBColor(0, 123, 255)  # #007bff
    COLOR_GRAY = RGBColor(108, 117, 125)  # #6c757d
    COLOR_DARK = RGBColor(52, 58, 64)  # #343a40
    
    # Typography standards
    FONT_NAME = "Calibri"
    FONT_SIZE_BODY = Pt(11)
    FONT_SIZE_H1 = Pt(18)
    FONT_SIZE_H2 = Pt(16)
    FONT_SIZE_H3 = Pt(14)
    FONT_SIZE_H4 = Pt(12)
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the document formatter.
        
        Args:
            template_path: Optional path to DOCX template file
        """
        if template_path:
            try:
                self.doc = Document(template_path)
                logger.info(f"📄 Loaded template from: {template_path}")
            except Exception as e:
                logger.warning(f"⚠️ Could not load template: {e}, using default")
                self.doc = Document()
        else:
            self.doc = Document()
        
        self._setup_styles()
        logger.info("✅ Professional document formatter initialized")
    
    def _setup_styles(self):
        """Set up document styles for consistent formatting."""
        try:
            styles = self.doc.styles
            
            # Configure Normal style
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = self.FONT_NAME
            normal_font.size = self.FONT_SIZE_BODY
            
            # Configure heading styles
            for level in range(1, 5):
                heading_style_name = f'Heading {level}'
                if heading_style_name in styles:
                    heading_style = styles[heading_style_name]
                    heading_font = heading_style.font
                    heading_font.name = self.FONT_NAME
                    heading_font.bold = True
                    
                    if level == 1:
                        heading_font.size = self.FONT_SIZE_H1
                        heading_font.color.rgb = self.COLOR_DARK
                    elif level == 2:
                        heading_font.size = self.FONT_SIZE_H2
                        heading_font.color.rgb = self.COLOR_DARK
                    elif level == 3:
                        heading_font.size = self.FONT_SIZE_H3
                    else:
                        heading_font.size = self.FONT_SIZE_H4
            
            logger.debug("✅ Document styles configured")
        except Exception as e:
            logger.warning(f"⚠️ Could not configure all styles: {e}")
    
    def create_cover_page(self, assessment_info: Dict[str, Any]):
        """
        Create professional cover page with assessment information.
        
        Args:
            assessment_info: Dictionary with title, date, architecture_name, etc.
        """
        logger.info("📄 Creating cover page")
        
        # Add title
        title = assessment_info.get("title", "AWS Well-Architected Framework Assessment Report")
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = self.COLOR_DARK
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(12)
        
        # Add architecture name if provided
        if assessment_info.get("architecture_name"):
            arch_para = self.doc.add_paragraph()
            arch_run = arch_para.add_run(assessment_info["architecture_name"])
            arch_run.font.size = Pt(16)
            arch_run.font.color.rgb = self.COLOR_GRAY
            arch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arch_para.space_after = Pt(24)
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Add assessment date
        date_str = assessment_info.get("assessment_date", datetime.now().strftime("%B %d, %Y"))
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(f"Assessment Date: {date_str}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add assessment ID if provided
        if assessment_info.get("assessment_id"):
            id_para = self.doc.add_paragraph()
            id_run = id_para.add_run(f"Assessment ID: {assessment_info['assessment_id']}")
            id_run.font.size = Pt(10)
            id_run.font.color.rgb = self.COLOR_GRAY
            id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break after cover
        self.doc.add_page_break()
        
        logger.info("✅ Cover page created")
    
    def create_table_of_contents(self, sections: List[str]):
        """
        Create table of contents with section links.
        
        Args:
            sections: List of section names
        """
        logger.info("📑 Creating table of contents")
        
        # Add TOC heading
        toc_heading = self.doc.add_heading("Table of Contents", level=1)
        toc_heading.space_after = Pt(12)
        
        # Add sections
        for idx, section in enumerate(sections, 1):
            para = self.doc.add_paragraph(style='List Number')
            para.add_run(section)
            para.space_after = Pt(6)
        
        # Page break after TOC
        self.doc.add_page_break()
        
        logger.info(f"✅ Table of contents created with {len(sections)} sections")
    
    def add_section_heading(self, text: str, level: int = 1):
        """
        Add a section heading with proper styling.
        
        Args:
            text: Heading text
            level: Heading level (1-4)
        """
        heading = self.doc.add_heading(text, level=level)
        heading.space_before = Pt(12)
        heading.space_after = Pt(6)
        return heading
    
    def add_paragraph(self, text: str, style: Optional[str] = None) -> Any:
        """
        Add a paragraph with proper styling.
        
        Args:
            text: Paragraph text
            style: Optional paragraph style
            
        Returns:
            Paragraph object
        """
        para = self.doc.add_paragraph(text, style=style)
        para.space_after = Pt(6)
        return para
    
    def add_bullet_list(self, items: List[str]):
        """
        Add a bulleted list.
        
        Args:
            items: List of items to add
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.space_after = Pt(3)
    
    def add_numbered_list(self, items: List[str]):
        """
        Add a numbered list.
        
        Args:
            items: List of items to add
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Number')
            para.space_after = Pt(3)
    
    def add_colored_text(self, text: str, color: RGBColor, bold: bool = False) -> Any:
        """
        Add text with specific color.
        
        Args:
            text: Text to add
            color: RGB color
            bold: Whether to make text bold
            
        Returns:
            Paragraph object
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.color.rgb = color
        if bold:
            run.font.bold = True
        return para
    
    def add_risk_indicator(self, risk_level: str) -> str:
        """
        Get text indicator for risk level.
        
        Args:
            risk_level: Risk level (critical, high, medium, low)
            
        Returns:
            Formatted risk indicator text
        """
        risk_indicators = {
            "critical": "🔴 CRITICAL",
            "high": "🔴 HIGH",
            "medium": "🟡 MEDIUM",
            "low": "🟢 LOW"
        }
        return risk_indicators.get(risk_level.lower(), "⚪ UNKNOWN")
    
    def add_priority_badge(self, priority: str) -> str:
        """
        Get text badge for priority level.
        
        Args:
            priority: Priority level (critical, high, medium, low)
            
        Returns:
            Formatted priority badge text
        """
        priority_badges = {
            "critical": "🔴 CRITICAL PRIORITY",
            "high": "🟠 HIGH PRIORITY",
            "medium": "🟡 MEDIUM PRIORITY",
            "low": "🟢 LOW PRIORITY"
        }
        return priority_badges.get(priority.lower(), "⚪ PRIORITY")
    
    def create_table(
        self,
        headers: List[str],
        rows: List[List[str]],
        col_widths: Optional[List[float]] = None
    ) -> Any:
        """
        Create a formatted table.
        
        Args:
            headers: List of column headers
            rows: List of row data (each row is a list of cell values)
            col_widths: Optional list of column widths in inches
            
        Returns:
            Table object
        """
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Set column widths if provided
        if col_widths:
            for idx, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[idx].width = Inches(width)
        
        # Add headers
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for row_idx, row_data in enumerate(rows, 1):
            cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                cells[col_idx].text = str(cell_data)
        
        # Add spacing after table
        self.doc.add_paragraph()
        
        return table
    
    def add_horizontal_line(self):
        """Add a horizontal line separator."""
        para = self.doc.add_paragraph()
        para.add_run("_" * 80)
        para.space_after = Pt(12)
    
    def add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()
    
    def add_hyperlink(self, paragraph: Any, text: str, url: str):
        """
        Add a hyperlink to a paragraph.
        
        Args:
            paragraph: Paragraph object to add link to
            text: Link text
            url: URL
        """
        # This is a workaround for adding hyperlinks in python-docx
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Add hyperlink styling
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
    
    def format_score_with_color(self, score: float) -> tuple:
        """
        Get color for score value.
        
        Args:
            score: Score value (0-100)
            
        Returns:
            Tuple of (color, risk_level)
        """
        if score >= 80:
            return (self.COLOR_GREEN, "low")
        elif score >= 60:
            return (self.COLOR_YELLOW, "medium")
        else:
            return (self.COLOR_RED, "high")
    
    def save_document(self, output_path: str) -> str:
        """
        Save the formatted document.
        
        Args:
            output_path: Path where document should be saved
            
        Returns:
            Path to saved document
        """
        try:
            self.doc.save(output_path)
            logger.info(f"✅ Document saved to: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"❌ Error saving document: {e}")
            raise
    
    def get_document(self) -> Document:
        """
        Get the underlying Document object.
        
        Returns:
            Document object
        """
        return self.doc
