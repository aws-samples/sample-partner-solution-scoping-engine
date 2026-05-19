"""DOCX File Validation and Recovery for WAFR Report Generation.

This module ensures DOCX files are valid and openable even when content is incomplete,
with error recovery mechanisms for generation failures.
"""

import logging
import os
import tempfile
import zipfile
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)


class DOCXValidationResult:
    """Result of DOCX file validation."""
    
    def __init__(
        self,
        is_valid: bool,
        file_path: str,
        errors: Optional[list] = None,
        warnings: Optional[list] = None,
        file_size: int = 0,
        can_open: bool = False
    ):
        """
        Initialize validation result.
        
        Args:
            is_valid: Whether the file is valid
            file_path: Path to the validated file
            errors: List of validation errors
            warnings: List of validation warnings
            file_size: Size of the file in bytes
            can_open: Whether the file can be opened
        """
        self.is_valid = is_valid
        self.file_path = file_path
        self.errors = errors or []
        self.warnings = warnings or []
        self.file_size = file_size
        self.can_open = can_open
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "is_valid": self.is_valid,
            "file_path": self.file_path,
            "errors": self.errors,
            "warnings": self.warnings,
            "file_size": self.file_size,
            "can_open": self.can_open
        }


class DOCXValidator:
    """Validator for DOCX files with recovery mechanisms."""
    
    def __init__(self):
        """Initialize the DOCX validator."""
        logger.info("📄 DOCX validator initialized")
    
    def validate_docx_file(self, file_path: str) -> DOCXValidationResult:
        """
        Validate a DOCX file for integrity and openability.
        
        Args:
            file_path: Path to the DOCX file to validate
            
        Returns:
            DOCXValidationResult with validation details
        """
        try:
            logger.info(f"🔍 Validating DOCX file: {file_path}")
            
            # Check if file exists
            if not os.path.exists(file_path):
                return DOCXValidationResult(
                    is_valid=False,
                    file_path=file_path,
                    errors=["File does not exist"],
                    can_open=False
                )
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            # Check if file is empty
            if file_size == 0:
                return DOCXValidationResult(
                    is_valid=False,
                    file_path=file_path,
                    errors=["File is empty"],
                    file_size=file_size,
                    can_open=False
                )
            
            # Validate DOCX structure (it's a ZIP file)
            errors = []
            warnings = []
            can_open = False
            
            try:
                # Check if it's a valid ZIP file
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Check for required DOCX components
                    required_files = [
                        'word/document.xml',
                        '[Content_Types].xml',
                        '_rels/.rels'
                    ]
                    
                    zip_contents = zip_file.namelist()
                    
                    for required_file in required_files:
                        if required_file not in zip_contents:
                            errors.append(f"Missing required file: {required_file}")
                    
                    # Check for corrupted files
                    try:
                        bad_files = zip_file.testzip()
                        if bad_files:
                            errors.append(f"Corrupted files detected: {bad_files}")
                    except Exception as test_error:
                        errors.append(f"Error testing ZIP integrity: {test_error}")
                    
                    # If no critical errors, file can likely be opened
                    if not errors:
                        can_open = True
                    elif len(errors) == 0 or all('Missing' not in e for e in errors):
                        # Minor errors, might still be openable
                        can_open = True
                        warnings.extend(errors)
                        errors = []
                
            except zipfile.BadZipFile:
                errors.append("File is not a valid ZIP/DOCX file")
            except Exception as zip_error:
                errors.append(f"Error validating ZIP structure: {zip_error}")
            
            # Try to open with python-docx if available
            try:
                from docx import Document
                doc = Document(file_path)
                # If we can create a Document object, it's openable
                can_open = True
                logger.info(f"✅ DOCX file validated successfully: {file_path}")
            except ImportError:
                # python-docx not available, rely on ZIP validation
                warnings.append("python-docx not available for full validation")
            except Exception as docx_error:
                errors.append(f"Cannot open with python-docx: {docx_error}")
                can_open = False
            
            is_valid = len(errors) == 0
            
            return DOCXValidationResult(
                is_valid=is_valid,
                file_path=file_path,
                errors=errors,
                warnings=warnings,
                file_size=file_size,
                can_open=can_open
            )
            
        except Exception as e:
            logger.error(f"❌ Error validating DOCX file: {e}")
            return DOCXValidationResult(
                is_valid=False,
                file_path=file_path,
                errors=[f"Validation error: {e}"],
                can_open=False
            )
    
    def recover_from_generation_failure(
        self,
        error: Exception,
        partial_content: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> Tuple[bool, str, str]:
        """
        Recover from DOCX generation failure by creating a minimal valid DOCX.
        
        Args:
            error: The error that occurred during generation
            partial_content: Any partial content that was generated
            output_path: Desired output path for the recovered file
            
        Returns:
            Tuple of (success, file_path, message)
        """
        try:
            logger.warning(f"⚠️ Attempting to recover from DOCX generation failure: {error}")
            
            # Create output path if not provided
            if not output_path:
                output_path = os.path.join(
                    tempfile.gettempdir(),
                    f"wafr_report_recovery_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                )
            
            # Try to create a minimal valid DOCX
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                
                doc = Document()
                
                # Add title
                title = doc.add_heading('WAFR Assessment Report - Partial', 0)
                title.alignment = 1  # Center
                
                # Add error notice
                error_para = doc.add_paragraph()
                error_run = error_para.add_run('⚠️ REPORT GENERATION ERROR')
                error_run.bold = True
                error_run.font.size = Pt(14)
                error_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
                error_para.alignment = 1  # Center
                
                # Add error details
                doc.add_paragraph()
                doc.add_paragraph(
                    f"An error occurred during report generation: {str(error)}"
                )
                doc.add_paragraph(
                    "This is a partial report with available information. "
                    "Please re-run the assessment for complete results."
                )
                
                # Add timestamp
                doc.add_paragraph()
                doc.add_paragraph(
                    f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                )
                
                # Add partial content if available
                if partial_content:
                    doc.add_page_break()
                    doc.add_heading('Partial Assessment Data', 1)
                    doc.add_paragraph(partial_content)
                
                # Add recovery notice
                doc.add_page_break()
                doc.add_heading('Recovery Information', 1)
                doc.add_paragraph(
                    "This report was automatically recovered after a generation failure. "
                    "The following actions are recommended:"
                )
                
                recommendations = [
                    "Review the error message above",
                    "Check that all required assessment data is available",
                    "Verify that all pillar assessments completed successfully",
                    "Re-run the assessment with complete data",
                    "Contact support if the issue persists"
                ]
                
                for rec in recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
                
                # Save the document
                doc.save(output_path)
                
                logger.info(f"✅ Recovery DOCX created successfully: {output_path}")
                
                return (
                    True,
                    output_path,
                    "Recovery DOCX created with error information and partial content"
                )
                
            except ImportError:
                # python-docx not available, create a text file instead
                logger.warning("⚠️ python-docx not available, creating text file")
                
                text_path = output_path.replace('.docx', '.txt')
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write("WAFR Assessment Report - Error Recovery\n")
                    f.write("=" * 50 + "\n\n")
                    f.write(f"Error: {str(error)}\n\n")
                    f.write("This is a recovery file created after report generation failed.\n")
                    f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                    
                    if partial_content:
                        f.write("\nPartial Content:\n")
                        f.write("-" * 50 + "\n")
                        f.write(partial_content)
                
                return (
                    True,
                    text_path,
                    "Recovery text file created (python-docx not available)"
                )
            
        except Exception as recovery_error:
            logger.error(f"❌ Recovery failed: {recovery_error}")
            return (
                False,
                "",
                f"Recovery failed: {recovery_error}"
            )
    
    def ensure_file_integrity(self, file_path: str) -> Tuple[bool, str]:
        """
        Ensure a DOCX file has integrity and can be opened.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (success, message)
        """
        try:
            logger.info(f"🔒 Ensuring file integrity: {file_path}")
            
            # Validate the file
            validation_result = self.validate_docx_file(file_path)
            
            if validation_result.is_valid and validation_result.can_open:
                logger.info(f"✅ File integrity confirmed: {file_path}")
                return (True, "File is valid and can be opened")
            
            # File has issues, try to repair
            if validation_result.can_open:
                # File can be opened despite warnings
                logger.warning(f"⚠️ File has warnings but can be opened: {file_path}")
                return (
                    True,
                    f"File can be opened with warnings: {', '.join(validation_result.warnings)}"
                )
            
            # File cannot be opened, attempt repair
            logger.warning(f"⚠️ File cannot be opened, attempting repair: {file_path}")
            
            # Try to repair by re-creating with python-docx
            try:
                from docx import Document
                
                # Try to open and re-save
                doc = Document(file_path)
                backup_path = file_path + '.backup'
                os.rename(file_path, backup_path)
                doc.save(file_path)
                
                # Validate repaired file
                repaired_validation = self.validate_docx_file(file_path)
                
                if repaired_validation.can_open:
                    logger.info(f"✅ File repaired successfully: {file_path}")
                    os.remove(backup_path)
                    return (True, "File was repaired and can now be opened")
                else:
                    # Repair failed, restore backup
                    os.remove(file_path)
                    os.rename(backup_path, file_path)
                    return (
                        False,
                        f"Repair failed: {', '.join(repaired_validation.errors)}"
                    )
                
            except ImportError:
                return (
                    False,
                    "Cannot repair file: python-docx not available"
                )
            except Exception as repair_error:
                return (
                    False,
                    f"Repair failed: {repair_error}"
                )
            
        except Exception as e:
            logger.error(f"❌ Error ensuring file integrity: {e}")
            return (False, f"Integrity check failed: {e}")
    
    def create_minimal_valid_docx(self, output_path: str, title: str = "WAFR Report") -> bool:
        """
        Create a minimal valid DOCX file.
        
        Args:
            output_path: Path where the file should be created
            title: Title for the document
            
        Returns:
            True if successful, False otherwise
        """
        try:
            logger.info(f"📄 Creating minimal valid DOCX: {output_path}")
            
            from docx import Document
            
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(
                f"This is a minimal valid DOCX file created on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            doc.save(output_path)
            
            logger.info(f"✅ Minimal DOCX created successfully: {output_path}")
            return True
            
        except ImportError:
            logger.error("❌ Cannot create minimal DOCX: python-docx not available")
            return False
        except Exception as e:
            logger.error(f"❌ Error creating minimal DOCX: {e}")
            return False
    
    def add_integrity_checks(self, file_path: str) -> Dict[str, Any]:
        """
        Add comprehensive integrity checks to a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with integrity check results
        """
        try:
            logger.info(f"🔍 Running integrity checks: {file_path}")
            
            checks = {
                "file_exists": os.path.exists(file_path),
                "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
                "is_zip": False,
                "has_required_files": False,
                "can_open_with_docx": False,
                "validation_result": None
            }
            
            if not checks["file_exists"]:
                return checks
            
            # Check if it's a valid ZIP
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    checks["is_zip"] = True
                    
                    # Check for required files
                    required_files = ['word/document.xml', '[Content_Types].xml']
                    zip_contents = zip_file.namelist()
                    checks["has_required_files"] = all(
                        f in zip_contents for f in required_files
                    )
            except Exception:
                pass
            
            # Try to open with python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                checks["can_open_with_docx"] = True
            except Exception:
                pass
            
            # Run full validation
            validation_result = self.validate_docx_file(file_path)
            checks["validation_result"] = validation_result.to_dict()
            
            logger.info(f"✅ Integrity checks completed: {file_path}")
            return checks
            
        except Exception as e:
            logger.error(f"❌ Error running integrity checks: {e}")
            return {
                "error": str(e),
                "file_exists": False
            }


# Singleton instance
_validator_instance = None


def get_docx_validator() -> DOCXValidator:
    """Get the singleton DOCX validator instance."""
    global _validator_instance
    if _validator_instance is None:
        _validator_instance = DOCXValidator()
    return _validator_instance
