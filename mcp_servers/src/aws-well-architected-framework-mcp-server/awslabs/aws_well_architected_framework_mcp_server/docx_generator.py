"""
ENTERPRISE-GRADE DOCX generation for WAFR reports.
Features: Professional styling, color preservation, rich formatting, branded templates.
"""

import base64
import io
import logging
import re
from datetime import datetime
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)

# AWS Brand Colors
AWS_NAVY = RGBColor(35, 47, 62)      # #232F3E
AWS_ORANGE = RGBColor(255, 153, 0)    # #FF9900
AWS_LIGHT_GRAY = RGBColor(248, 249, 250)  # #F8F9FA
AWS_DARK_GRAY = RGBColor(108, 117, 125)   # #6C757D

# Score Colors
SCORE_EXCELLENT = RGBColor(40, 167, 69)    # #28A745 - Green
SCORE_GOOD = RGBColor(92, 184, 92)         # #5CB85C - Light Green  
SCORE_FAIR = RGBColor(255, 193, 7)         # #FFC107 - Yellow
SCORE_NEEDS_IMPROVEMENT = RGBColor(253, 126, 20)  # #FD7E14 - Orange
SCORE_CRITICAL = RGBColor(220, 53, 69)     # #DC3545 - Red

# Risk Colors
RISK_HIGH = RGBColor(220, 53, 69)          # #DC3545 - Red
RISK_MEDIUM = RGBColor(255, 193, 7)        # #FFC107 - Yellow
RISK_LOW = RGBColor(40, 167, 69)           # #28A745 - Green


def create_professional_title_page(doc: Document, assessment_data: Dict[str, Any]) -> None:
    """Create enterprise-grade title page with AWS branding."""
    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AWS Well-Architected Framework")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = AWS_NAVY
    title_run.bold = True
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Assessment Report")
    subtitle_run.font.name = 'Segoe UI'
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = AWS_ORANGE
    subtitle_run.bold = True
    
    # Architecture name (if available)
    if assessment_data and assessment_data.get('document_analysis', {}).get('architectural_patterns'):
        doc.add_paragraph()  # Spacing
        arch_para = doc.add_paragraph()
        arch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        arch_run = arch_para.add_run("Architecture Assessment")
        arch_run.font.name = 'Segoe UI'
        arch_run.font.size = Pt(18)
        arch_run.font.color.rgb = AWS_NAVY
    
    # Assessment date
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # More spacing
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Assessment Date: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.name = 'Segoe UI'
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = AWS_DARK_GRAY
    
    # Page break
    doc.add_page_break()


def add_branded_header_footer(doc: Document) -> None:
    """Add professional AWS-branded headers and footers."""
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run("AWS Well-Architected Framework Assessment")
    header_run.font.name = 'Segoe UI'
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = AWS_DARK_GRAY
    
    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')} | AWS Well-Architected Framework Tool")
    footer_run.font.name = 'Segoe UI'
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = AWS_DARK_GRAY


def get_score_color(score: float) -> RGBColor:
    """Get appropriate color for score based on value."""
    if score >= 90:
        return SCORE_EXCELLENT
    elif score >= 80:
        return SCORE_GOOD
    elif score >= 70:
        return SCORE_FAIR
    elif score >= 60:
        return SCORE_NEEDS_IMPROVEMENT
    else:
        return SCORE_CRITICAL


def get_risk_color(risk_level: str) -> RGBColor:
    """Get appropriate color for risk level."""
    risk_level = risk_level.lower()
    if 'high' in risk_level or 'critical' in risk_level:
        return RISK_HIGH
    elif 'medium' in risk_level:
        return RISK_MEDIUM
    else:
        return RISK_LOW


def create_professional_table(doc: Document, headers: List[str], rows: List[List[str]], 
                            title: str = None, score_column: int = None) -> None:
    """Create enterprise-grade table with professional styling."""
    if title:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = 'Segoe UI'
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = AWS_NAVY
        title_run.bold = True
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        
        # Header styling
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_run = cell_para.runs[0]
        cell_run.font.name = 'Segoe UI'
        cell_run.font.size = Pt(11)
        cell_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        cell_run.bold = True
        
        # Header background color (AWS Navy) - using direct XML manipulation
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '232F3E')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception:
            # Fallback: just use bold formatting if XML manipulation fails
            pass
    
    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(cell_data)
            
            # Cell styling
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_run = cell_para.runs[0]
            cell_run.font.name = 'Segoe UI'
            cell_run.font.size = Pt(10)
            
            # Color-code score columns
            if score_column is not None and i == score_column:
                try:
                    score_value = float(re.findall(r'\d+\.?\d*', str(cell_data))[0])
                    cell_run.font.color.rgb = get_score_color(score_value)
                    cell_run.bold = True
                except (ValueError, IndexError):
                    cell_run.font.color.rgb = AWS_DARK_GRAY
            else:
                cell_run.font.color.rgb = AWS_DARK_GRAY
    
    # Table borders and spacing
    table.style = 'Table Grid'
    doc.add_paragraph()  # Spacing after table


def strip_css_and_scripts(html_content: str) -> str:
    """Remove CSS style blocks and script tags from HTML while preserving class info."""
    if not html_content:
        return ""
    
    # Remove style blocks but preserve class attributes for color mapping
    html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove script blocks
    html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    
    return html_content


def extract_table_data_with_styling(table_html: str) -> Dict[str, Any]:
    """Extract table data from HTML table with styling information."""
    headers = []
    rows = []
    table_title = None
    
    # Extract table title if present (look for h3/h4 before table)
    title_match = re.search(r'<h[34][^>]*>([^<]+)</h[34]>', table_html, re.IGNORECASE)
    if title_match:
        table_title = title_match.group(1).strip()
    
    # Find header row
    header_pattern = r'<thead[^>]*>.*?<tr[^>]*>(.*?)</tr>.*?</thead>'
    header_match = re.search(header_pattern, table_html, re.DOTALL | re.IGNORECASE)
    if header_match:
        header_content = header_match.group(1)
        header_cell_pattern = r'<th[^>]*>(.*?)</th>'
        for cell_match in re.finditer(header_cell_pattern, header_content, re.DOTALL | re.IGNORECASE):
            cell_text = re.sub(r'<[^>]+>', '', cell_match.group(1)).strip()
            headers.append(cell_text)
    
    # Find data rows
    tbody_pattern = r'<tbody[^>]*>(.*?)</tbody>'
    tbody_match = re.search(tbody_pattern, table_html, re.DOTALL | re.IGNORECASE)
    if tbody_match:
        tbody_content = tbody_match.group(1)
    else:
        tbody_content = table_html  # Fallback to entire table
    
    row_pattern = r'<tr[^>]*>(.*?)</tr>'
    for row_match in re.finditer(row_pattern, tbody_content, re.DOTALL | re.IGNORECASE):
        row_content = row_match.group(1)
        
        # Skip header rows in tbody
        if '<th' in row_content:
            continue
        
        # Find all cells (td)
        cell_pattern = r'<td[^>]*>(.*?)</td>'
        cells = []
        for cell_match in re.finditer(cell_pattern, row_content, re.DOTALL | re.IGNORECASE):
            cell_content = cell_match.group(1)
            
            # Extract text and preserve important formatting
            cell_text = cell_content
            
            # Handle strong/bold text
            cell_text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'\1', cell_text, flags=re.DOTALL | re.IGNORECASE)
            cell_text = re.sub(r'<b[^>]*>(.*?)</b>', r'\1', cell_text, flags=re.DOTALL | re.IGNORECASE)
            
            # Clean up remaining HTML tags
            cell_text = re.sub(r'<[^>]+>', '', cell_text).strip()
            
            # Clean up HTML entities
            cell_text = cell_text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
            cells.append(cell_text)
        
        if cells:  # Only add non-empty rows
            rows.append(cells)
    
    # Detect score column (look for % symbols)
    score_column = None
    for i, header in enumerate(headers):
        if any(keyword in header.lower() for keyword in ['score', 'confidence', 'percentage', '%']):
            score_column = i
            break
    
    # If no header match, check data for % symbols
    if score_column is None and rows:
        for i in range(len(rows[0])):
            if any('%' in str(row[i]) for row in rows):
                score_column = i
                break
    
    return {
        'title': table_title,
        'headers': headers,
        'rows': rows,
        'score_column': score_column
    }


def add_formatted_paragraph(doc: Document, text: str, style_class: str = None) -> None:
    """Add paragraph with enterprise formatting based on content and style class."""
    if not text.strip():
        return
    
    # ENTERPRISE FIX: Detect sub-headings that should be bold (e.g., "Excellent (90-100%)", "High Priority", etc.)
    # These are paragraphs that contain ONLY bold text and should be formatted as sub-headings
    subheading_pattern = r'^<p[^>]*>\s*<strong[^>]*>([^<]+)</strong>\s*</p>$'
    subheading_match = re.match(subheading_pattern, text.strip(), re.IGNORECASE)
    
    if subheading_match:
        # This is a sub-heading - format as bold paragraph with larger font
        para = doc.add_paragraph()
        heading_text = subheading_match.group(1).strip()
        run = para.add_run(heading_text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(12)  # Slightly larger for sub-headings
        run.font.color.rgb = AWS_NAVY
        run.bold = True
        return
    
    # ENTERPRISE FIX: Detect capability sub-headings with percentages (e.g., "Infrastructure as Code: 95%")
    capability_pattern = r'<strong[^>]*>([^<]+:\s*\d+%)</strong>'
    if re.search(capability_pattern, text, re.IGNORECASE):
        para = doc.add_paragraph()
        # Extract the capability name and percentage
        cap_match = re.search(capability_pattern, text, re.IGNORECASE)
        if cap_match:
            run = para.add_run(cap_match.group(1))
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            run.font.color.rgb = AWS_NAVY
            run.bold = True
        return
    
    # ENTERPRISE FIX: Detect service layer headings (e.g., "Compute Layer (EC2/Auto Scaling)")
    layer_pattern = r'<strong[^>]*>([^<]*Layer[^<]*)</strong>'
    if re.search(layer_pattern, text, re.IGNORECASE):
        para = doc.add_paragraph()
        layer_match = re.search(layer_pattern, text, re.IGNORECASE)
        if layer_match:
            run = para.add_run(layer_match.group(1))
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            run.font.color.rgb = AWS_NAVY
            run.bold = True
        return
    
    # ENTERPRISE FIX: Detect "Strengths:" and "Recommendations:" labels
    label_pattern = r'<strong[^>]*>(Strengths|Recommendations|Phase \d+)[^<]*</strong>'
    if re.search(label_pattern, text, re.IGNORECASE):
        para = doc.add_paragraph()
        label_match = re.search(label_pattern, text, re.IGNORECASE)
        if label_match:
            run = para.add_run(label_match.group(0).replace('<strong>', '').replace('</strong>', '').strip())
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            run.font.color.rgb = AWS_NAVY
            run.bold = True
        return
    
    para = doc.add_paragraph()
    
    # Detect and handle score badges
    score_badge_pattern = r'<span class="score-badge[^"]*">([^<]+)</span>'
    if re.search(score_badge_pattern, text):
        # Handle score badge formatting
        parts = re.split(score_badge_pattern, text)
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                if part.strip():
                    run = para.add_run(re.sub(r'<[^>]+>', '', part))
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(11)
                    run.font.color.rgb = AWS_DARK_GRAY
            else:  # Score badge content
                run = para.add_run(part)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                run.bold = True
                
                # Extract score for color coding
                try:
                    score_match = re.search(r'(\d+\.?\d*)%', part)
                    if score_match:
                        score = float(score_match.group(1))
                        run.font.color.rgb = get_score_color(score)
                    else:
                        run.font.color.rgb = AWS_ORANGE
                except (ValueError, AttributeError):
                    run.font.color.rgb = AWS_ORANGE
        return
    
    # Handle risk level formatting
    if any(risk in text.lower() for risk in ['high risk', 'medium risk', 'low risk', 'critical']):
        run = para.add_run(re.sub(r'<[^>]+>', '', text))
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.bold = True
        
        if 'high risk' in text.lower() or 'critical' in text.lower():
            run.font.color.rgb = RISK_HIGH
        elif 'medium risk' in text.lower():
            run.font.color.rgb = RISK_MEDIUM
        else:
            run.font.color.rgb = RISK_LOW
        return
    
    # Handle bullet points
    if text.strip().startswith('•') or text.strip().startswith('-'):
        para.style = 'List Bullet'
        text = text.strip()[1:].strip()  # Remove bullet
    
    # Regular paragraph formatting
    clean_text = re.sub(r'<[^>]+>', '', text)
    
    # Handle bold text
    bold_pattern = r'<(?:strong|b)[^>]*>(.*?)</(?:strong|b)>'
    if re.search(bold_pattern, text, re.IGNORECASE):
        parts = re.split(bold_pattern, text, flags=re.IGNORECASE)
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                if part.strip():
                    clean_part = re.sub(r'<[^>]+>', '', part)
                    run = para.add_run(clean_part)
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(11)
                    run.font.color.rgb = AWS_DARK_GRAY
            else:  # Bold text
                run = para.add_run(part)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                run.font.color.rgb = AWS_NAVY
                run.bold = True
    else:
        # Simple paragraph
        run = para.add_run(clean_text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = AWS_DARK_GRAY


def process_list_with_nesting(doc: Document, list_html: str, indent_level: int = 0) -> None:
    """
    Process HTML lists with proper nesting, bold text handling, and indentation.
    Handles nested <ul> and <ol> elements with proper hierarchy.
    """
    # Find all top-level list items
    # Use a more careful approach to handle nested lists
    li_pattern = r'<li[^>]*>(.*?)</li>'
    
    # First, let's process the list structure more carefully
    items = []
    current_pos = 0
    list_content = list_html
    
    # Find each <li> and its content, being careful about nesting
    while True:
        li_start = list_content.find('<li', current_pos)
        if li_start == -1:
            break
        
        # Find the matching </li>
        li_content_start = list_content.find('>', li_start) + 1
        
        # Count nested <li> tags to find the correct closing </li>
        depth = 1
        pos = li_content_start
        while depth > 0 and pos < len(list_content):
            next_open = list_content.find('<li', pos)
            next_close = list_content.find('</li>', pos)
            
            if next_close == -1:
                break
            
            if next_open != -1 and next_open < next_close:
                depth += 1
                pos = next_open + 3
            else:
                depth -= 1
                if depth == 0:
                    li_content_end = next_close
                pos = next_close + 5
        
        if depth == 0:
            item_content = list_content[li_content_start:li_content_end]
            items.append(item_content)
            current_pos = li_content_end + 5
        else:
            break
    
    for item_content in items:
        # Check if this item has nested lists
        has_nested_ul = '<ul' in item_content
        has_nested_ol = '<ol' in item_content
        
        if has_nested_ul or has_nested_ol:
            # Extract the text before the nested list
            if has_nested_ul:
                nested_start = item_content.find('<ul')
            else:
                nested_start = item_content.find('<ol')
            
            main_text = item_content[:nested_start].strip()
            nested_list = item_content[nested_start:]
            
            # Add the main item
            if main_text:
                add_list_item_with_formatting(doc, main_text, indent_level)
            
            # Process nested list with increased indent
            process_list_with_nesting(doc, nested_list, indent_level + 1)
        else:
            # Simple item without nesting
            add_list_item_with_formatting(doc, item_content, indent_level)


def add_list_item_with_formatting(doc: Document, item_html: str, indent_level: int = 0) -> None:
    """
    Add a list item with proper formatting, handling bold text and visual indicators.
    """
    # Clean the HTML but preserve bold markers
    clean_text = re.sub(r'<[^>]+>', '', item_html).strip()
    
    if not clean_text:
        return
    
    # Create paragraph with appropriate style
    para = doc.add_paragraph(style='List Bullet')
    
    # Set indentation based on level
    if indent_level > 0:
        para.paragraph_format.left_indent = Inches(0.25 * indent_level)
    
    # Check for bold patterns in original HTML
    bold_match = re.search(r'<strong[^>]*>(.*?)</strong>', item_html, re.IGNORECASE)
    
    if bold_match:
        # Has bold text - split and format
        bold_text = re.sub(r'<[^>]+>', '', bold_match.group(1)).strip()
        
        # Find position of bold text in clean text
        bold_pos = clean_text.find(bold_text)
        
        if bold_pos >= 0:
            # Text before bold
            if bold_pos > 0:
                before_run = para.add_run(clean_text[:bold_pos])
                before_run.font.name = 'Segoe UI'
                before_run.font.size = Pt(11)
                before_run.font.color.rgb = AWS_DARK_GRAY
            
            # Bold text
            bold_run = para.add_run(bold_text)
            bold_run.font.name = 'Segoe UI'
            bold_run.font.size = Pt(11)
            bold_run.font.color.rgb = AWS_DARK_GRAY
            bold_run.bold = True
            
            # Text after bold
            after_pos = bold_pos + len(bold_text)
            if after_pos < len(clean_text):
                after_run = para.add_run(clean_text[after_pos:])
                after_run.font.name = 'Segoe UI'
                after_run.font.size = Pt(11)
                after_run.font.color.rgb = AWS_DARK_GRAY
        else:
            # Fallback - just add the text
            run = para.add_run(clean_text)
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            run.font.color.rgb = AWS_DARK_GRAY
    else:
        # No bold - simple text
        run = para.add_run(clean_text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = AWS_DARK_GRAY


def extract_html_elements_with_nesting(html_content: str) -> List[str]:
    """
    Extract HTML elements while properly handling nested lists.
    Returns a list of complete HTML elements in document order.
    """
    elements = []
    pos = 0
    
    while pos < len(html_content):
        # Find the next tag
        tag_start = html_content.find('<', pos)
        if tag_start == -1:
            break
        
        # Check what type of tag this is
        tag_match = re.match(r'<(h[1-6]|p|ul|ol|table)[\s>]', html_content[tag_start:], re.IGNORECASE)
        
        if tag_match:
            tag_name = tag_match.group(1).lower()
            
            # Find the matching closing tag, accounting for nesting
            if tag_name in ('ul', 'ol'):
                # Handle nested lists
                element_end = find_matching_close_tag(html_content, tag_start, tag_name)
            else:
                # Simple close tag search for non-nested elements
                close_tag = f'</{tag_name}>'
                close_pos = html_content.find(close_tag, tag_start)
                if close_pos != -1:
                    element_end = close_pos + len(close_tag)
                else:
                    pos = tag_start + 1
                    continue
            
            if element_end > tag_start:
                element = html_content[tag_start:element_end]
                elements.append(element)
                pos = element_end
            else:
                pos = tag_start + 1
        else:
            pos = tag_start + 1
    
    return elements


def find_matching_close_tag(html_content: str, start_pos: int, tag_name: str) -> int:
    """
    Find the matching closing tag for a potentially nested element.
    Handles nested <ul> and <ol> tags properly.
    """
    open_tag = f'<{tag_name}'
    close_tag = f'</{tag_name}>'
    
    depth = 0
    pos = start_pos
    
    while pos < len(html_content):
        next_open = html_content.find(open_tag, pos if depth == 0 else pos + 1)
        next_close = html_content.find(close_tag, pos + 1)
        
        if next_close == -1:
            return -1
        
        if depth == 0:
            # First opening tag
            depth = 1
            pos = start_pos + len(open_tag)
            continue
        
        if next_open != -1 and next_open < next_close:
            # Found another opening tag before closing
            depth += 1
            pos = next_open + len(open_tag)
        else:
            # Found closing tag
            depth -= 1
            if depth == 0:
                return next_close + len(close_tag)
            pos = next_close + len(close_tag)
    
    return -1


def process_html_content_sequentially(doc: Document, html_content: str, processed_headings: set) -> None:
    """
    CRITICAL FIX: Process HTML content in sequential order to maintain proper document structure.
    This ensures tables appear within their pillar sections, not grouped together at the start.
    """
    if not html_content:
        return
    
    # Extract elements manually to handle nested lists properly
    elements = extract_html_elements_with_nesting(html_content)
    
    logger.info(f"📄 Processing {len(elements)} HTML elements sequentially")
    
    for element in elements:
        element = element.strip()
        if not element:
            continue
        
        # Determine element type
        if re.match(r'<h1[^>]*>', element, re.IGNORECASE):
            # H1 - Major section heading
            heading_text = re.sub(r'<[^>]+>', '', element).strip()
            if heading_text and heading_text not in processed_headings:
                processed_headings.add(heading_text)
                heading = doc.add_heading(heading_text, level=1)
                if heading.runs:
                    heading.runs[0].font.name = 'Segoe UI'
                    heading.runs[0].font.size = Pt(18)
                    heading.runs[0].font.color.rgb = AWS_NAVY
                    heading.runs[0].bold = True
        
        elif re.match(r'<h2[^>]*>', element, re.IGNORECASE):
            # H2 - Section heading
            heading_text = re.sub(r'<[^>]+>', '', element).strip()
            if heading_text and heading_text not in processed_headings:
                processed_headings.add(heading_text)
                heading = doc.add_heading(heading_text, level=1)
                if heading.runs:
                    heading.runs[0].font.name = 'Segoe UI'
                    heading.runs[0].font.size = Pt(16)
                    heading.runs[0].font.color.rgb = AWS_NAVY
                    heading.runs[0].bold = True
        
        elif re.match(r'<h[34][^>]*>', element, re.IGNORECASE):
            # H3/H4 - Subsection heading
            heading_text = re.sub(r'<[^>]+>', '', element).strip()
            if heading_text and heading_text not in processed_headings:
                processed_headings.add(heading_text)
                level = 3 if '<h3' in element.lower() else 3
                heading = doc.add_heading(heading_text, level=level)
                if heading.runs:
                    heading.runs[0].font.name = 'Segoe UI'
                    heading.runs[0].font.size = Pt(14)
                    heading.runs[0].font.color.rgb = AWS_ORANGE
                    heading.runs[0].bold = True
        
        elif re.match(r'<h[56][^>]*>', element, re.IGNORECASE):
            # H5/H6 - Minor heading
            heading_text = re.sub(r'<[^>]+>', '', element).strip()
            if heading_text and heading_text not in processed_headings:
                processed_headings.add(heading_text)
                heading = doc.add_heading(heading_text, level=3)
                if heading.runs:
                    heading.runs[0].font.name = 'Segoe UI'
                    heading.runs[0].font.size = Pt(12)
                    heading.runs[0].font.color.rgb = AWS_DARK_GRAY
                    heading.runs[0].bold = True
        
        elif re.match(r'<table[^>]*>', element, re.IGNORECASE):
            # TABLE - Process inline where it appears
            table_data = extract_table_data_with_styling(element)
            if table_data['headers'] and table_data['rows']:
                create_professional_table(
                    doc,
                    table_data['headers'],
                    table_data['rows'],
                    title=table_data['title'],
                    score_column=table_data['score_column']
                )
        
        elif re.match(r'<p[^>]*>', element, re.IGNORECASE):
            # Paragraph
            add_formatted_paragraph(doc, element, '')
        
        elif re.match(r'<[uo]l[^>]*>', element, re.IGNORECASE):
            # List (ordered or unordered) - handle nested lists properly
            process_list_with_nesting(doc, element)


def html_to_structured_content(html_content: str) -> List[Dict[str, Any]]:
    """Convert HTML to structured content preserving formatting information."""
    if not html_content:
        return []
    
    content_blocks = []
    
    # Split by major elements
    elements = re.findall(r'<(h[1-6]|p|ul|ol|table)[^>]*>.*?</\1>|<(h[1-6]|p)[^>]*>.*?(?=<(?:h[1-6]|p|ul|ol|table|div)|$)', 
                         html_content, re.DOTALL | re.IGNORECASE)
    
    for element_match in re.finditer(r'<(h[1-6]|p|ul|ol|table)[^>]*>(.*?)</\1>|<(h[1-6]|p)[^>]*>([^<]*)', 
                                   html_content, re.DOTALL | re.IGNORECASE):
        full_match = element_match.group(0)
        tag = element_match.group(1) or element_match.group(3)
        content = element_match.group(2) or element_match.group(4)
        
        if not tag or not content or not content.strip():
            continue
        
        content_blocks.append({
            'type': tag.lower(),
            'content': content.strip(),
            'full_html': full_match,
            'classes': re.findall(r'class="([^"]*)"', full_match)
        })
    
    return content_blocks


async def generate_wafr_docx(html_content: str, output_path: str, assessment_data: Dict[str, Any] = None) -> str:
    """
    Generate ENTERPRISE-GRADE WAFR DOCX with professional styling and rich formatting.
    
    Features:
    - Professional title page with AWS branding
    - Color-coded score badges and risk levels
    - Enterprise-grade tables with professional styling
    - Branded headers and footers
    - Rich content sections with business impact
    - Implementation roadmaps and recommendations
    
    Args:
        html_content: HTML content to convert to DOCX
        output_path: Path where the DOCX should be saved
        assessment_data: Assessment data for enhanced content generation
        
    Returns:
        Path to the generated DOCX file
    """
    try:
        logger.info(f"🏢 Generating ENTERPRISE-GRADE WAFR DOCX at {output_path}")
        logger.info(f"📄 HTML content length: {len(html_content)}")
        
        # Create DOCX document with enterprise settings
        doc = Document()
        
        # Set professional margins (1.25" as per enterprise standards)
        for section in doc.sections:
            section.top_margin = Inches(1.25)
            section.bottom_margin = Inches(1.25)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 🎨 ENTERPRISE FEATURE 1: Professional Title Page
        logger.info("✅ Creating professional title page with AWS branding")
        create_professional_title_page(doc, assessment_data)
        
        # 🎨 ENTERPRISE FEATURE 2: Branded Headers and Footers
        logger.info("✅ Adding branded headers and footers")
        add_branded_header_footer(doc)
        
        # Process HTML content with enterprise formatting
        clean_html = strip_css_and_scripts(html_content)
        logger.info(f"📝 Cleaned HTML length: {len(clean_html)}")
        
        processed_headings = set()  # Track to avoid duplicates
        
        # CRITICAL FIX: Process content in sequential order to maintain proper structure
        # This ensures tables appear within their pillar sections, not grouped together
        process_html_content_sequentially(doc, clean_html, processed_headings)
        
        # 🎨 ENTERPRISE FEATURE 7: Business Impact Section - DISABLED (redundant with Risk Analysis section)
        # if assessment_data and assessment_data.get('pillar_assessments'):
        #     logger.info("💼 Adding business impact analysis section")
        #     add_business_impact_section(doc, assessment_data)
        
        # 🎨 ENTERPRISE FEATURE 8: Implementation Roadmap - DISABLED to avoid duplicate with template
        # The HTML template already includes a comprehensive roadmap section
        # if assessment_data and assessment_data.get('high_priority_actions'):
        #     logger.info("🗺️ Adding implementation roadmap section")
        #     add_implementation_roadmap(doc, assessment_data)
        
        # Save document
        doc.save(output_path)
        logger.info(f"✅ Successfully generated ENTERPRISE-GRADE WAFR DOCX: {output_path}")
        logger.info(f"📊 Document features: Professional title page, branded headers/footers, color-coded content, enterprise tables")
        
        return output_path
        
    except Exception as e:
        logger.error(f"❌ Error generating ENTERPRISE WAFR DOCX: {e}", exc_info=True)
        raise Exception(f"Failed to generate ENTERPRISE WAFR DOCX: {e}")


def add_business_impact_section(doc: Document, assessment_data: Dict[str, Any]) -> None:
    """Add enterprise-grade business impact analysis section."""
    doc.add_heading("Business Impact Analysis", level=1)
    
    # Overall business impact
    overall_score = assessment_data.get('overall_score', 0)
    
    para = doc.add_paragraph()
    run = para.add_run("Overall Business Risk Level: ")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(12)
    run.font.color.rgb = AWS_NAVY
    run.bold = True
    
    if overall_score >= 80:
        risk_text = "LOW RISK - Architecture demonstrates strong alignment with AWS best practices"
        risk_color = RISK_LOW
    elif overall_score >= 60:
        risk_text = "MEDIUM RISK - Some areas require attention to optimize business outcomes"
        risk_color = RISK_MEDIUM
    else:
        risk_text = "HIGH RISK - Critical improvements needed to ensure business continuity"
        risk_color = RISK_HIGH
    
    risk_run = para.add_run(risk_text)
    risk_run.font.name = 'Segoe UI'
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = risk_color
    risk_run.bold = True
    
    # Cost impact analysis
    doc.add_paragraph()
    cost_para = doc.add_paragraph()
    cost_run = cost_para.add_run("💰 Cost Impact: ")
    cost_run.font.name = 'Segoe UI'
    cost_run.font.size = Pt(11)
    cost_run.font.color.rgb = AWS_NAVY
    cost_run.bold = True
    
    cost_score = assessment_data.get('pillar_assessments', {}).get('cost_optimization', {}).get('score', 0)
    if cost_score < 70:
        cost_impact = "Significant cost optimization opportunities identified. Potential savings of 20-40% through recommended improvements."
    else:
        cost_impact = "Cost optimization is well-managed with minor improvement opportunities."
    
    cost_text_run = cost_para.add_run(cost_impact)
    cost_text_run.font.name = 'Segoe UI'
    cost_text_run.font.size = Pt(11)
    cost_text_run.font.color.rgb = AWS_DARK_GRAY


def add_implementation_roadmap(doc: Document, assessment_data: Dict[str, Any]) -> None:
    """Add enterprise-grade implementation roadmap section with actionable remediation steps."""
    # Import remediation guidance
    try:
        from .core.remediation_guidance import get_remediation_guidance
    except ImportError:
        def get_remediation_guidance(cap): return {}
    
    doc.add_heading("Implementation Roadmap", level=1)
    
    high_priority = assessment_data.get('high_priority_actions', [])
    
    if high_priority:
        # Phase 1: Critical Issues (0-30 days)
        doc.add_heading("Phase 1: Critical Security & Reliability (0-30 days)", level=2)
        
        critical_actions = [action for action in high_priority if action.get('priority') == 'critical']
        high_actions = [action for action in high_priority if action.get('priority') == 'high']
        phase1_actions = critical_actions + high_actions[:2]
        
        if phase1_actions:
            for action in phase1_actions[:3]:  # Top 3 critical/high
                # Add action title
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(f"{action.get('title', 'Critical Action')}")
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                run.font.color.rgb = RISK_HIGH
                run.bold = True
                
                # Get remediation guidance for this action
                cap_name = action.get('title', '').lower().replace('implement ', '').replace('enhance ', '').replace(' ', '_')
                guidance = get_remediation_guidance(cap_name)
                
                if guidance:
                    # Add AWS services to use
                    services = guidance.get('aws_services', [])[:4]
                    if services:
                        svc_para = doc.add_paragraph(style='List Bullet 2')
                        svc_run = svc_para.add_run(f"AWS Services: {', '.join(services)}")
                        svc_run.font.name = 'Segoe UI'
                        svc_run.font.size = Pt(10)
                        svc_run.font.color.rgb = AWS_NAVY
                    
                    # Add implementation steps
                    steps = guidance.get('implementation_steps', [])[:3]
                    for step in steps:
                        step_para = doc.add_paragraph(style='List Bullet 2')
                        step_run = step_para.add_run(f"→ {step}")
                        step_run.font.name = 'Segoe UI'
                        step_run.font.size = Pt(10)
                        step_run.font.color.rgb = AWS_DARK_GRAY
                    
                    # Add effort estimate
                    effort = guidance.get('estimated_effort', '2-4 weeks')
                    effort_para = doc.add_paragraph(style='List Bullet 2')
                    effort_run = effort_para.add_run(f"Estimated Effort: {effort}")
                    effort_run.font.name = 'Segoe UI'
                    effort_run.font.size = Pt(10)
                    effort_run.font.italic = True
        
        # Phase 2: Performance & Cost Optimization (30-90 days)
        doc.add_heading("Phase 2: Performance & Cost Optimization (30-90 days)", level=2)
        
        medium_actions = [action for action in high_priority if action.get('priority') == 'medium']
        if medium_actions:
            for action in medium_actions[:3]:  # Top 3 medium
                # Add action title
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(f"{action.get('title', 'Optimization Action')}")
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                run.font.color.rgb = RISK_MEDIUM
                run.bold = True
                
                # Get remediation guidance
                cap_name = action.get('title', '').lower().replace('implement ', '').replace('enhance ', '').replace(' ', '_')
                guidance = get_remediation_guidance(cap_name)
                
                if guidance:
                    # Add AWS services
                    services = guidance.get('aws_services', [])[:4]
                    if services:
                        svc_para = doc.add_paragraph(style='List Bullet 2')
                        svc_run = svc_para.add_run(f"AWS Services: {', '.join(services)}")
                        svc_run.font.name = 'Segoe UI'
                        svc_run.font.size = Pt(10)
                        svc_run.font.color.rgb = AWS_NAVY
                    
                    # Add top 2 implementation steps
                    steps = guidance.get('implementation_steps', [])[:2]
                    for step in steps:
                        step_para = doc.add_paragraph(style='List Bullet 2')
                        step_run = step_para.add_run(f"→ {step}")
                        step_run.font.name = 'Segoe UI'
                        step_run.font.size = Pt(10)
                        step_run.font.color.rgb = AWS_DARK_GRAY
        
        # Success metrics
        doc.add_heading("Success Metrics", level=3)
        metrics_para = doc.add_paragraph()
        metrics_run = metrics_para.add_run("📊 Target Improvements:")
        metrics_run.font.name = 'Segoe UI'
        metrics_run.font.size = Pt(11)
        metrics_run.font.color.rgb = AWS_NAVY
        metrics_run.bold = True
        
        doc.add_paragraph("• Overall WAFR Score: Target 85%+ within 90 days", style='List Bullet')
        doc.add_paragraph("• Security Score: Target 80%+ within 30 days", style='List Bullet')
        doc.add_paragraph("• Cost Optimization: Target 15-25% cost reduction", style='List Bullet')
