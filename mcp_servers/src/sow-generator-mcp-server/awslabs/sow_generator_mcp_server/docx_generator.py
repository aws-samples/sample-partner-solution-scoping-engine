# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License"). You may not use this file except in compliance
# with the License. A copy of the License is located at
#
#    http://www.apache.org/licenses/LICENSE-2.0
#
# or in the 'license' file accompanying this file. This file is distributed on an 'AS IS' BASIS, WITHOUT WARRANTIES
# OR CONDITIONS OF ANY KIND, express or implied. See the License for the specific language governing permissions
# and limitations under the License.

"""DOCX generation module using python-docx for SOW documents."""

import base64
import io
import logging
import re
import tempfile
from typing import Optional, List, Dict, Any
from html.parser import HTMLParser

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


logger = logging.getLogger(__name__)


def strip_css_and_scripts(html_content: str) -> str:
    """Remove CSS style blocks and script tags from HTML."""
    if not html_content:
        return ""
    
    # Remove style blocks
    html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove script blocks
    html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove inline style attributes
    html_content = re.sub(r'\s+style\s*=\s*["\'][^"\']*["\']', '', html_content, flags=re.IGNORECASE)
    
    # Remove CSS @page rules that might be in the content
    html_content = re.sub(r'@page\s*{[^}]*}', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    
    return html_content


def extract_images_from_html(html_content: str) -> List[Dict[str, Any]]:
    """Extract base64 images from HTML content."""
    images = []
    img_pattern = r'<img[^>]+src="data:image/([^;]+);base64,([^"]+)"[^>]*>'
    
    for match in re.finditer(img_pattern, html_content):
        img_format = match.group(1)
        img_data = match.group(2)
        
        try:
            # Decode base64 image
            img_bytes = base64.b64decode(img_data)
            images.append({
                'format': img_format,
                'data': img_bytes,
                'placeholder': match.group(0)
            })
        except Exception as e:
            logger.warning(f"Failed to decode image: {e}")
    
    return images


def extract_table_data(table_html: str) -> List[List[str]]:
    """Extract table data from HTML table."""
    rows = []
    
    # Find all table rows
    row_pattern = r'<tr[^>]*>(.*?)</tr>'
    for row_match in re.finditer(row_pattern, table_html, re.DOTALL | re.IGNORECASE):
        row_content = row_match.group(1)
        
        # Find all cells (th or td)
        cell_pattern = r'<(?:th|td)[^>]*>(.*?)</(?:th|td)>'
        cells = []
        for cell_match in re.finditer(cell_pattern, row_content, re.DOTALL | re.IGNORECASE):
            cell_content = cell_match.group(1)
            # Strip HTML tags and clean up text
            cell_text = re.sub(r'<[^>]+>', '', cell_content).strip()
            # Clean up HTML entities
            cell_text = cell_text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
            cells.append(cell_text)
        
        if cells:  # Only add non-empty rows
            rows.append(cells)
    
    return rows


def html_to_text(html_content: str) -> str:
    """Convert HTML to plain text while preserving structure."""
    if not html_content:
        return ""
    
    # Replace common HTML elements with text equivalents
    text = html_content
    
    # Convert headings
    text = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'\n\1\n', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Convert paragraphs
    text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Convert line breaks
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    
    # Convert list items
    text = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove list containers
    text = re.sub(r'</?(?:ul|ol)[^>]*>', '', text, flags=re.IGNORECASE)
    
    # Convert strong/bold
    text = re.sub(r'<(?:strong|b)[^>]*>(.*?)</(?:strong|b)>', r'\1', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Convert emphasis/italic  
    text = re.sub(r'<(?:em|i)[^>]*>(.*?)</(?:em|i)>', r'\1', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Clean up whitespace
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Clean up HTML entities
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&nbsp;', ' ')
    
    return text.strip()


async def generate_sow_docx(html_content: str, output_path: str, sow_request=None) -> str:
    """Generate DOCX from HTML content using python-docx.
    
    Args:
        html_content: HTML content to convert to DOCX
        output_path: Path where the DOCX should be saved
        sow_request: Original SOW request object with data (unused)
        
    Returns:
        Path to the generated DOCX file
        
    Raises:
        Exception: If DOCX generation fails
    """
    try:
        logger.info(f"Generating DOCX with python-docx at {output_path}")
        logger.info(f"HTML content length: {len(html_content)}")
        
        # Strip CSS and scripts first
        clean_html = strip_css_and_scripts(html_content)
        logger.info(f"Cleaned HTML length: {len(clean_html)}")
        
        # Create DOCX document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('Statement of Work', 0)
        title_run = title.runs[0] if title.runs else title.add_run('Statement of Work')
        title_run.font.color.rgb = RGBColor(35, 47, 62)
        
        # Split content into sections
        sections_content = re.split(r'<h[1-3][^>]*>', clean_html, flags=re.IGNORECASE)
        headings = re.findall(r'<h[1-3][^>]*>(.*?)</h[1-3]>', clean_html, flags=re.DOTALL | re.IGNORECASE)
        
        # Process each section
        for i, section_content in enumerate(sections_content[1:], 0):  # Skip first empty section
            if i < len(headings):
                # Add heading
                heading_text = re.sub(r'<[^>]+>', '', headings[i]).strip()
                if heading_text:
                    heading = doc.add_heading(heading_text, level=2)
                    heading_run = heading.runs[0] if heading.runs else heading.add_run(heading_text)
                    heading_run.font.color.rgb = RGBColor(35, 47, 62)
            
            # Extract and process tables separately
            table_pattern = r'<table[^>]*>.*?</table>'
            tables = re.findall(table_pattern, section_content, re.DOTALL | re.IGNORECASE)
            
            # Remove tables from section content for text processing
            section_text = re.sub(table_pattern, '[TABLE_PLACEHOLDER]', section_content, flags=re.DOTALL | re.IGNORECASE)
            
            # Convert remaining content to text
            text_content = html_to_text(section_text)
            
            # Process text content line by line
            table_index = 0
            for line in text_content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                    
                if line == '[TABLE_PLACEHOLDER]' and table_index < len(tables):
                    # Insert table
                    table_data = extract_table_data(tables[table_index])
                    if table_data:
                        # Determine number of columns from first row
                        max_cols = max(len(row) for row in table_data) if table_data else 1
                        
                        # Create table
                        table = doc.add_table(rows=len(table_data), cols=max_cols)
                        table.style = 'Table Grid'
                        
                        # Populate table
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < max_cols:
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = cell_data
                                    
                                    # Style header row
                                    if row_idx == 0:
                                        cell.paragraphs[0].runs[0].bold = True
                    
                    table_index += 1
                    
                elif line.startswith('•'):
                    # Bullet point
                    doc.add_paragraph(line[1:].strip(), style='List Bullet')
                else:
                    # Regular paragraph
                    paragraph = doc.add_paragraph(line)
        
        # Handle appendices
        if 'appendices-section' in clean_html.lower():
            # Add page break for appendices
            doc.add_page_break()
            
            # Add appendices title
            appendices_title = doc.add_heading('Appendices', level=1)
            appendices_title.runs[0].font.color.rgb = RGBColor(35, 47, 62)
            
            # Extract appendices content
            appendices_match = re.search(r'<div[^>]*class="[^"]*appendices-section[^"]*"[^>]*>(.*?)</div>\s*$', clean_html, re.DOTALL | re.IGNORECASE)
            if not appendices_match:
                # Try alternative extraction - find the appendices section and extract until end
                start_match = re.search(r'<div[^>]*class="[^"]*appendices-section[^"]*"[^>]*>', clean_html, re.IGNORECASE)
                if start_match:
                    start_pos = start_match.end()
                    # Find the last </div> and extract content
                    content_after = clean_html[start_pos:]
                    last_div = content_after.rfind('</div>')
                    if last_div > 0:
                        appendices_content = content_after[:last_div]
                        appendices_match = type('obj', (object,), {'group': lambda self, n: appendices_content if n == 1 else None})()
            if appendices_match:
                appendices_content = appendices_match.group(1)
                
                # Split into individual appendices
                appendix_sections = re.split(r'<div[^>]*class="[^"]*appendix-section[^"]*"[^>]*>', appendices_content, flags=re.IGNORECASE)
                
                for i, appendix_content in enumerate(appendix_sections[1:], 1):
                    # Add page break before each appendix (except first)
                    if i > 1:
                        doc.add_page_break()
                    
                    # Extract appendix title
                    title_match = re.search(r'<h2[^>]*>(.*?)</h2>', appendix_content, re.IGNORECASE)
                    if title_match:
                        title_text = re.sub(r'<[^>]+>', '', title_match.group(1)).strip()
                        appendix_heading = doc.add_heading(title_text, level=2)
                        appendix_heading.runs[0].font.color.rgb = RGBColor(35, 47, 62)
                    
                    # Extract and add description
                    desc_match = re.search(r'<p[^>]*>(.*?)</p>', appendix_content, re.IGNORECASE)
                    if desc_match:
                        desc_text = re.sub(r'<[^>]+>', '', desc_match.group(1)).strip()
                        if desc_text:
                            doc.add_paragraph(desc_text)
                    
                    # Extract and add images FIRST (for diagram appendix)
                    images = extract_images_from_html(appendix_content)
                    for img in images:
                        try:
                            # Create temporary file for image
                            temp_file = tempfile.NamedTemporaryFile(suffix=f'.{img["format"]}', delete=False)
                            temp_file.write(img['data'])
                            temp_file.close()
                            
                            # Add image to document
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run()
                            run.add_picture(temp_file.name, width=Inches(6))
                            
                            # Clean up temp file
                            import os
                            os.unlink(temp_file.name)
                            
                        except Exception as e:
                            logger.warning(f"Failed to add image to DOCX: {e}")
                            doc.add_paragraph('[Architecture Diagram could not be displayed]')
                    
                    # Process markdown content in appendices
                    markdown_content_match = re.search(r'<div[^>]*class="[^"]*markdown-content[^"]*"[^>]*>(.*?)</div>', appendix_content, re.DOTALL | re.IGNORECASE)
                    if markdown_content_match:
                        markdown_html = markdown_content_match.group(1)
                        
                        # Process tables in markdown content
                        table_pattern = r'<table[^>]*>.*?</table>'
                        tables = re.findall(table_pattern, markdown_html, re.DOTALL | re.IGNORECASE)
                        
                        # Remove tables for text processing
                        text_html = re.sub(table_pattern, '[TABLE_PLACEHOLDER]', markdown_html, flags=re.DOTALL | re.IGNORECASE)
                        
                        # Convert remaining HTML to text and process
                        text_content = html_to_text(text_html)
                        table_index = 0
                        
                        for line in text_content.split('\n'):
                            line = line.strip()
                            if not line:
                                continue
                            
                            if line == '[TABLE_PLACEHOLDER]' and table_index < len(tables):
                                # Insert table
                                table_data = extract_table_data(tables[table_index])
                                if table_data:
                                    # Determine number of columns
                                    max_cols = max(len(row) for row in table_data) if table_data else 1
                                    
                                    # Create table
                                    table = doc.add_table(rows=len(table_data), cols=max_cols)
                                    table.style = 'Table Grid'
                                    
                                    # Populate table
                                    for row_idx, row_data in enumerate(table_data):
                                        for col_idx, cell_data in enumerate(row_data):
                                            if col_idx < max_cols:
                                                cell = table.cell(row_idx, col_idx)
                                                cell.text = cell_data
                                                
                                                # Style header row
                                                if row_idx == 0:
                                                    for run in cell.paragraphs[0].runs:
                                                        run.bold = True
                                
                                table_index += 1
                                
                            elif line.startswith('•'):
                                # Bullet point
                                doc.add_paragraph(line[1:].strip(), style='List Bullet')
                            elif line.startswith('#'):
                                # Heading
                                heading_level = len(line) - len(line.lstrip('#'))
                                heading_text = line.lstrip('#').strip()
                                if heading_text:
                                    heading = doc.add_heading(heading_text, level=min(heading_level + 2, 6))
                                    heading.runs[0].font.color.rgb = RGBColor(35, 47, 62)
                            else:
                                # Regular paragraph
                                if line:
                                    doc.add_paragraph(line)
        
        # Save document
        doc.save(output_path)
        
        logger.info(f"Successfully generated DOCX: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Failed to generate DOCX with python-docx: {e}", exc_info=True)
        raise Exception(f"DOCX generation failed: {str(e)}")